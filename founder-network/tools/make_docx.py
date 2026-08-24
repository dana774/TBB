from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = "/tmp/claude-0/-home-user-TBB/ab62350a-c683-5f9e-8199-204331d5ec27/scratchpad/out/"
INK = RGBColor(0x1A, 0x2B, 0x3C)
ACCENT = RGBColor(0xB8, 0x76, 0x3E)
MUTED = RGBColor(0x5A, 0x6B, 0x7C)

d = Document()
for s in d.sections:
    s.left_margin = s.right_margin = Inches(0.9)
    s.top_margin = Inches(0.8); s.bottom_margin = Inches(0.75)

st = d.styles["Normal"]
st.font.name = "Calibri"; st.font.size = Pt(10.5); st.font.color.rgb = INK
st.paragraph_format.space_after = Pt(6); st.paragraph_format.line_spacing = 1.15


def kicker(t):
    p = d.add_paragraph(); r = p.add_run(t.upper())
    r.font.size = Pt(8.5); r.font.bold = True; r.font.color.rgb = ACCENT
    p.paragraph_format.space_after = Pt(2)


def h1(t):
    p = d.add_paragraph(); r = p.add_run(t)
    r.font.size = Pt(19); r.font.bold = True; r.font.color.rgb = INK
    p.paragraph_format.space_after = Pt(8)


def h2(t):
    p = d.add_paragraph(); r = p.add_run(t)
    r.font.size = Pt(12); r.font.bold = True; r.font.color.rgb = INK
    p.paragraph_format.space_before = Pt(14); p.paragraph_format.space_after = Pt(4)


def body(t, italic=False, small=False):
    p = d.add_paragraph(); r = p.add_run(t)
    r.font.italic = italic
    if small:
        r.font.size = Pt(8); r.font.color.rgb = MUTED
    return p


def bullet(t):
    p = d.add_paragraph(t, style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    for r in p.runs:
        r.font.size = Pt(10.5); r.font.color.rgb = INK


def fillable(label, lines=3, width=6.6):
    p = d.add_paragraph(); r = p.add_run(label)
    r.font.bold = True; r.font.size = Pt(10)
    p.paragraph_format.space_after = Pt(3)
    t = d.add_table(rows=lines, cols=1)
    t.style = "Table Grid"; t.alignment = WD_TABLE_ALIGNMENT.LEFT
    for row in t.rows:
        row.height = Pt(20)
        row.cells[0].width = Inches(width)
    d.add_paragraph().paragraph_format.space_after = Pt(4)


def table(headers, rows, widths):
    t = d.add_table(rows=1, cols=len(headers)); t.style = "Table Grid"
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = ""
        run = c.paragraphs[0].add_run(h)
        run.font.bold = True; run.font.size = Pt(9); run.font.color.rgb = INK
        c.width = Inches(widths[i])
        shd = OxmlElement("w:shd"); shd.set(qn("w:fill"), "F2F5F8")
        c._tc.get_or_add_tcPr().append(shd)
    for r in rows:
        cells = t.add_row().cells
        for i, v in enumerate(r):
            cells[i].text = ""
            run = cells[i].paragraphs[0].add_run(v)
            run.font.size = Pt(9); run.font.color.rgb = INK
            cells[i].width = Inches(widths[i])
    d.add_paragraph().paragraph_format.space_after = Pt(4)


kicker("ACC-02 · Accelerator")
h1("Cohort Workbook & Session Templates")
body("The Brand Blueprint is your founder-facing ecosystem, Value Growth Partners is the advisory "
     "and execution engine behind it, and Growth OS turns plans into operating infrastructure — "
     "this workbook is what you fill in as a cohort runs, so the sessions leave artifacts behind "
     "instead of notes you never reopen.")

h2("How to use this workbook")
bullet("Duplicate the session template for each session. Seven modules, seven copies.")
bullet("Complete the pre-work section before the session, not during it. Sessions start from your "
       "real figures and lose their value when they start from guesses.")
bullet("Fill the commitment block before you leave the room, while the decision is still fresh.")
bullet("Bring the previous session's commitment block to the next session. That is the only "
       "accountability mechanism the cohort has.")

d.add_page_break()

kicker("Template · duplicate per session")
h1("Session record")
table(["Field", "Entry"],
      [["Module", ""], ["Session date", ""], ["Facilitator", ""], ["Attending from my team", ""]],
      [1.8, 4.8])

h2("Pre-work — complete before the session")
fillable("The figures I am bringing (actual, not estimated):")
fillable("The decision I want to leave this session having made:")
fillable("What I tried since the last session, and what happened:", lines=3)

h2("In session — working notes")
fillable("What I learned that contradicts what I believed coming in:", lines=3)
fillable("Peer feedback worth keeping (who said it, and what):", lines=3)

h2("Artifact produced")
body("Each module is meant to produce something you keep. Name it, say where it lives, and who "
     "else has seen it.")
table(["Artifact", "Where it lives", "Reviewed by", "Date"],
      [["", "", "", ""], ["", "", "", ""]],
      [2.4, 2.0, 1.4, 0.8])

h2("Commitment block — complete before leaving")
table(["Commitment", "Owner", "By when", "What done looks like"],
      [["", "", "", ""], ["", "", "", ""], ["", "", "", ""]],
      [2.3, 1.1, 1.0, 2.2])
body("If you cannot state what done looks like, the commitment is not yet a commitment.", small=True)

d.add_page_break()

kicker("Template")
h1("Peer review sheet")
body("Used when you read another founder's work. Be specific and be kind — in that order. Vague "
     "encouragement wastes the reviewer's seat.")
table(["Prompt", "Response"],
      [["Whose work am I reviewing?", ""],
       ["What is clearly true and well evidenced here?", ""],
       ["Where does the reasoning depend on an assumption that is not stated?", ""],
       ["What would I need to see to believe the conclusion?", ""],
       ["One thing I would do differently, and why", ""]],
      [2.6, 4.0])

h2("Confidentiality")
body("What founders share in cohort sessions stays in the room. Do not repeat another founder's "
     "figures, customers, suppliers, or plans outside the cohort, during or after it.")

d.add_page_break()

kicker("Template")
h1("Module close-out")
body("Complete at the end of each module, before the next begins.")
table(["Question", "Response"],
      [["What changed in the business as a result of this module?", ""],
       ["What did I commit to that I did not do?", ""],
       ["What is now blocking me that the cohort cannot resolve?", ""],
       ["Which library collection or partner should that route to?", ""]],
      [3.2, 3.4])

h2("Carry-forward")
fillable("Open items carried into the next module:", lines=3)

h2("Notices")
body("© 2026 Value Growth Partners. Provided to Founder Network members for their own use. Not "
     "legal, financial, investment, or tax advice. Do not redistribute, resell, or share outside "
     "your organization.", small=True)
body("Educational content only. Participation in a cohort does not guarantee funding, buyer "
     "placement, partnership, or any commercial result; outcomes depend on your own execution, "
     "market, and circumstances. Cohort structure and session dates may change — the schedule "
     "issued to your cohort supersedes this workbook. Notes you record here may include other "
     "founders' confidential information; store and dispose of this workbook accordingly.",
     small=True)
body("Founder Network member resource · The Brand Blueprint · Powered by Value Growth Partners",
     italic=True, small=True)

d.save(OUT + "BB_Accelerator_Cohort-Workbook-Session-Templates_v1.docx")
print("ACC-02 ok")
